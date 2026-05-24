import streamlit as st
import pandas as pd
from sqlalchemy import create_engine
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

# --- KONEKSI DATABASE ---
DB_URL = st.secrets["DB_URL"]
engine = create_engine(DB_URL, pool_size=5, max_overflow=10)

def format_rupiah(x):
    try: return f"{float(x):,.0f}".replace(',', '.')
    except (ValueError, TypeError): return x

def get_vol_sat_combined(v1, s1, v2, s2):
    v1_str = str(v1).replace(".0", "") if pd.notna(v1) else "0"
    s1_str = str(s1).strip() if pd.notna(s1) else ""
    v2_str = str(v2).replace(".0", "") if pd.notna(v2) else "0"
    s2_str = str(s2).strip() if pd.notna(s2) else ""
    if s2_str in ["", "-"] or v2_str == "0" or v2_str == "": return f"{v1_str} {s1_str}"
    return f"{v1_str} {s1_str} x {v2_str} {s2_str}"

@st.cache_data(ttl=60)
def load_active_rab():
    """Hanya mengambil RAB yang berstatus AKTIF (Is_Active = 1)"""
    conn = engine.connect()
    try:
        df_utama = pd.read_sql("SELECT * FROM rab_utama WHERE \"Is_Active\" = 1", conn)
        if not df_utama.empty:
            ids = tuple(df_utama['ID_RAB'].tolist())
            if len(ids) == 1:
                df_detail = pd.read_sql(f"SELECT * FROM rab_detail WHERE \"ID_RAB\" = '{ids[0]}'", conn)
            else:
                df_detail = pd.read_sql(f"SELECT * FROM rab_detail WHERE \"ID_RAB\" IN {ids}", conn)
        else:
            df_detail = pd.DataFrame()
    except Exception as e:
        df_utama, df_detail = pd.DataFrame(), pd.DataFrame()
    conn.close()
    return df_utama, df_detail

# --- FUNGSI AI GEMINI ---
def generate_narasi_tor(kegiatan, total_anggaran, sasaran, list_belanja, poin_tambahan):
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-3.1-flash-lite')
        
        prompt = f"""
        Anda adalah perencana anggaran ahli di Fakultas Ilmu Budaya Universitas Mulawarman. 
        Tugas Anda adalah menulis Bab Pendahuluan untuk Term of Reference (TOR) kegiatan.
        
        DATA KEGIATAN:
        - Nama Kegiatan: {kegiatan}
        - Sasaran: {sasaran}
        - Total Anggaran: Rp {total_anggaran}
        - Item Utama yang dibeli: {list_belanja}
        - Catatan Panitia: {poin_tambahan}
        
        INSTRUKSI PENULISAN:
        1. Tulis dengan gaya bahasa Indonesia formal, birokratis, dan akademis.
        2. Format output harus rapi, HANYA terdiri dari 2 bagian utama: "A. Latar Belakang" dan "B. Tujuan Kegiatan".
        3. Jelaskan mengapa item belanja tersebut penting untuk mendukung sasaran kegiatan.
        4. Jangan gunakan kata "Saya", "Kami", atau format markdown tebal (**) yang berlebihan. Gunakan teks biasa.
        """
        
        respons = model.generate_content(prompt)
        return respons.text
    except Exception as e:
        return f"[ERROR AI] Gagal menyusun narasi. Pastikan GEMINI_API_KEY di Streamlit Secrets sudah benar. Detail: {e}"

# --- FUNGSI BUILDER MICROSOFT WORD (.DOCX) ---
def build_docx(kegiatan, tahun, ketua, tgl, lokasi, narasi, df_det_keg, total_biaya):
    doc = Document()
    
    # KOP & JUDUL
    p_judul = doc.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_judul = p_judul.add_run("KERANGKA ACUAN KERJA (TERM OF REFERENCE)\n")
    r_judul.bold = True
    r_judul.font.size = Pt(14)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"FAKULTAS ILMU BUDAYA - UNIVERSITAS MULAWARMAN\nTAHUN ANGGARAN {tahun}")
    r_sub.bold = True
    
    doc.add_paragraph("_" * 70) # Garis Pemisah
    
    # METADATA
    doc.add_paragraph(f"Nama Kegiatan\t\t: {kegiatan}")
    doc.add_paragraph(f"Penanggung Jawab\t: {ketua}")
    doc.add_paragraph(f"Waktu Pelaksanaan\t: {tgl}")
    doc.add_paragraph(f"Tempat Pelaksanaan\t: {lokasi}")
    
    # NARASI AI
    doc.add_paragraph("\n" + narasi)
    
    # TABEL RINCIAN RAB
    doc.add_paragraph("\nC. Rincian Anggaran Belanja (RAB)").bold = True
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'No'
    hdr[1].text = 'Uraian Belanja'
    hdr[2].text = 'Volume'
    hdr[3].text = 'Jumlah Biaya (Rp)'
    
    for i, row in enumerate(df_det_keg.itertuples(), 1):
        r_cells = table.add_row().cells
        r_cells[0].text = str(i)
        r_cells[1].text = row.Uraian
        r_cells[2].text = get_vol_sat_combined(row.Vol_1, row.Sat_1, row.Vol_2, row.Sat_2)
        r_cells[3].text = format_rupiah(row.Total_Biaya)
        
    tot_row = table.add_row().cells
    tot_row[1].text = "TOTAL KESELURUHAN"
    tot_row[3].text = format_rupiah(total_biaya)
    
    # TANDA TANGAN
    doc.add_paragraph("\n\n")
    p_ttd = doc.add_paragraph(f"Samarinda, ........................ {tahun}\nPenanggung Jawab Kegiatan\n\n\n\n")
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ttd.add_run(f"({ketua})").bold = True
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# =====================================================================
# TAMPILAN HALAMAN UTAMA
# =====================================================================
def show_page():
    st.title("🤖 Asisten Penyusun TOR Otomatis")
    st.caption("Didukung oleh Google Gemini AI. Data anggaran terintegrasi langsung dengan versi RAB Aktif terkini.")
    
    df_utama, df_detail = load_active_rab()
    
    if df_utama.empty:
        st.warning("⚠️ Belum ada dokumen RAB yang berstatus AKTIF. Silakan aktifkan RAB di modul Pengolah RAB terlebih dahulu.")
        return

    # Inisialisasi State agar teks AI tidak hilang saat tombol ditekan
    if 'draft_ai' not in st.session_state:
        st.session_state.draft_ai = ""
    
    kegiatan_list = sorted(df_utama['Kegiatan'].unique())
    keg_display = {k: k.title() for k in kegiatan_list}
    
    tab_setup, tab_drafting, tab_cetak = st.tabs(["1️⃣ Pilih & Setup Kegiatan", "2️⃣ Drafting AI Gemini", "3️⃣ Finalisasi & Cetak Word"])

    # --- TAB 1: SETUP ---
    with tab_setup:
        st.subheader("Data Dasar Kegiatan")
        pilih_keg = st.selectbox("Pilih Kegiatan (Dari RAB Aktif):", kegiatan_list, format_func=lambda x: keg_display[x])
        
        # Tarik data RAB
        df_keg_utama = df_utama[df_utama['Kegiatan'] == pilih_keg].iloc[0]
        df_keg_det = df_detail[df_detail['ID_RAB'] == df_keg_utama['ID_RAB']]
        tot_rp = df_keg_det['Total_Biaya'].sum()
        
        st.info(f"**Sasaran RAB:** {df_keg_utama['Sasaran']}\n\n**Total Pagu:** Rp {format_rupiah(tot_rp)}")
        
        st.markdown("---")
        st.subheader("Lengkapi Informasi Pelaksanaan")
        col1, col2 = st.columns(2)
        in_ketua = col1.text_input("Nama Penanggung Jawab / Ketua", placeholder="Budi Santoso, S.Hum., M.A.")
        in_lokasi = col2.text_input("Lokasi Pelaksanaan", placeholder="Ruang Serbaguna FIB")
        in_tgl = col1.text_input("Waktu Pelaksanaan", placeholder="12 - 14 Agustus 2027")
        in_poin = st.text_area("Ide / Catatan Bebas untuk AI (Opsional)", placeholder="Ketik ide singkat Anda di sini. Misal: 'Alat sudah rusak 5 tahun, penting untuk media promosi penerimaan mahasiswa baru'. AI akan merangkainya menjadi paragraf formal.")

    # --- TAB 2: AI DRAFTING ---
    with tab_drafting:
        st.subheader("Penyusunan Narasi Berbantu AI")
        st.markdown("AI Gemini akan membaca rincian barang yang dibeli dan ide bebas Anda, lalu menyusunnya menjadi Bab Latar Belakang dan Tujuan berbahasa baku.")
        
        if st.button("✨ Hasilkan Narasi dengan Gemini", type="primary"):
            with st.spinner("Gemini sedang berpikir dan mengetik... (Bisa memakan waktu 5-10 detik)"):
                list_barang = ", ".join(df_keg_det['Uraian'].tolist()[:10]) # Ambil 10 barang utama agar prompt tidak kepanjangan
                if len(df_keg_det) > 10: list_barang += ", dll."
                
                hasil = generate_narasi_tor(
                    kegiatan=pilih_keg.title(),
                    total_anggaran=format_rupiah(tot_rp),
                    sasaran=df_keg_utama['Sasaran'],
                    list_belanja=list_barang,
                    poin_tambahan=in_poin
                )
                st.session_state.draft_ai = hasil
                st.success("Draft berhasil dibuat!")
                
        if st.session_state.draft_ai:
            st.markdown("### Hasil Ketikan AI (Bisa Diedit Manual)")
            teks_final = st.text_area("Silakan perbaiki kata-kata yang kurang pas di bawah ini:", value=st.session_state.draft_ai, height=400)
            st.session_state.teks_final_tor = teks_final # Simpan untuk dicetak

    # --- TAB 3: CETAK ---
    with tab_cetak:
        st.subheader("Cetak Dokumen Resmi")
        if 'teks_final_tor' in st.session_state and st.session_state.teks_final_tor:
            st.success("Narasi telah siap. Data RAB akan dilampirkan secara otomatis ke dalam dokumen Word.")
            
            file_word = build_docx(
                kegiatan=pilih_keg.title(),
                tahun=df_keg_utama['Tahun'],
                ketua=in_ketua,
                tgl=in_tgl,
                lokasi=in_lokasi,
                narasi=st.session_state.teks_final_tor,
                df_det_keg=df_keg_det,
                total_biaya=tot_rp
            )
            
            st.download_button(
                label="📥 Download TOR (.docx)",
                data=file_word,
                file_name=f"TOR_{pilih_keg.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        else:
            st.warning("⚠️ Anda belum menyusun narasi. Silakan proses di tab '2️⃣ Drafting AI Gemini' terlebih dahulu.")
